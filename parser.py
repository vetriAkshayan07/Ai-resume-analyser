import re
import pdfplumber
import docx
import os


# ---------------------------------------------------------------------------
# Common skills keyword list (expandable)
# ---------------------------------------------------------------------------
SKILLS_KEYWORDS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "php",
    "swift", "kotlin", "go", "rust", "scala", "r", "matlab",
    "html", "css", "react", "angular", "vue", "node", "django", "flask",
    "spring", "laravel", "express", "fastapi", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "sql", "mysql", "postgresql", "mongodb",
    "redis", "elasticsearch", "aws", "azure", "gcp", "docker", "kubernetes",
    "git", "linux", "machine learning", "deep learning", "nlp",
    "data analysis", "data science", "power bi", "tableau", "excel", "hadoop",
    "spark", "kafka", "rest api", "graphql", "agile", "scrum", "devops",
    "ci/cd", "jenkins", "terraform", "ansible", "figma", "photoshop",
    "communication", "leadership", "teamwork", "problem solving",
    "project management", "time management",
]

# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def extract_text_from_pdf(filepath: str) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        text = f"[PDF extraction error: {e}]"
    return text.strip()


def extract_text_from_docx(filepath: str) -> str:
    """Extract all text from a DOCX file using python-docx."""
    text = ""
    try:
        document = docx.Document(filepath)
        for para in document.paragraphs:
            text += para.text + "\n"
        # Also extract tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        text = f"[DOCX extraction error: {e}]"
    return text.strip()


def extract_text(filepath: str) -> str:
    """Dispatch extraction based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    return ""


# ---------------------------------------------------------------------------
# Field extraction helpers
# ---------------------------------------------------------------------------

def extract_email(text: str) -> str:
    """Extract the first email address found in text."""
    pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    match = re.search(pattern, text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    """Extract the first phone number found in text."""
    pattern = r"(\+?\d[\d\s\-().]{7,}\d)"
    match = re.search(pattern, text)
    return match.group(0).strip() if match else ""


def extract_name(text: str) -> str:
    """
    Heuristic: The candidate name is usually in the first 3 non-empty lines.
    We pick the first line that looks like a name (2-4 words, no digits, no @).
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:6]:
        # Skip lines with email, phone, URLs, or very long strings
        if "@" in line or re.search(r"\d{3,}", line) or len(line) > 60:
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
            return line
    return lines[0] if lines else "Unknown"


def extract_skills(text: str) -> list:
    """Return a list of skills found in the resume text (case-insensitive match)."""
    text_lower = text.lower()
    found = []
    for skill in SKILLS_KEYWORDS:
        if re.search(r"\b" + re.escape(skill) + r"\b", text_lower):
            found.append(skill.title())
    return list(dict.fromkeys(found))  # preserve order, remove duplicates


def extract_education(text: str) -> str:
    """Extract education-related sentences."""
    education_keywords = [
        "bachelor", "master", "phd", "doctorate", "b.sc", "m.sc", "b.tech",
        "m.tech", "mba", "b.e", "m.e", "diploma", "degree", "university",
        "college", "institute", "school", "graduation", "graduated",
    ]
    lines = text.splitlines()
    edu_lines = []
    capture = False
    for line in lines:
        lower = line.lower()
        if "education" in lower or "academic" in lower or "qualification" in lower:
            capture = True
        if capture and line.strip():
            edu_lines.append(line.strip())
            if len(edu_lines) > 8:
                break
        elif any(kw in lower for kw in education_keywords):
            edu_lines.append(line.strip())
    return " | ".join(edu_lines[:5]) if edu_lines else "Not specified"


def extract_experience(text: str) -> str:
    """
    Extract years of experience from explicit mentions like '3 years', '5+ years',
    or from work experience section headers.
    """
    # Try to find explicit year mentions
    year_pattern = r"(\d+\+?\s*years?\s*(?:of\s*)?(?:experience|exp)?)"
    matches = re.findall(year_pattern, text, re.IGNORECASE)
    if matches:
        return matches[0].strip()

    # Fallback: detect work history section and count date ranges
    experience_section = _extract_section(text, ["experience", "employment", "work history"])
    if experience_section:
        date_pattern = r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*[\s,]+\d{4})"
        dates = re.findall(date_pattern, experience_section, re.IGNORECASE)
        if dates:
            return f"Found {len(dates) // 2 or 1}+ positions"
    return "Not specified"


def extract_certifications(text: str) -> str:
    """Extract certifications from the resume."""
    cert_keywords = [
        "certified", "certification", "certificate", "aws certified",
        "pmp", "cissp", "ceh", "ccna", "comptia", "google certified",
        "microsoft certified", "oracle certified", "scrum master", "csm",
    ]
    lines = text.splitlines()
    cert_lines = []
    for line in lines:
        lower = line.lower()
        if any(kw in lower for kw in cert_keywords):
            cert_lines.append(line.strip())
    cert_section = _extract_section(text, ["certifications", "certificates", "credentials"])
    if cert_section:
        cert_lines.extend(cert_section.splitlines())
    unique = list(dict.fromkeys([l for l in cert_lines if l]))
    return " | ".join(unique[:5]) if unique else "None listed"


def extract_location(text: str) -> str:
    """Attempt to extract location / city from top of resume."""
    # Common country / city pattern after name/email/phone block
    location_pattern = r"\b([A-Z][a-z]+(?:[\s,]+[A-Z][a-z]+)*,?\s*(?:[A-Z]{2}|\w+))\b"
    lines = text.splitlines()[:15]
    for line in lines:
        if re.search(r"\b(city|state|location|address|india|usa|uk|canada|australia)\b", line, re.IGNORECASE):
            return line.strip()
        # Simple heuristic: short line with comma
        if "," in line and 3 < len(line) < 60 and not "@" in line:
            return line.strip()
    return "Not specified"


def _extract_section(text: str, section_names: list) -> str:
    """
    Extract a block of text following a section header matching any name in section_names.
    Returns up to 20 lines after the header.
    """
    lines = text.splitlines()
    capture = False
    captured = []
    for i, line in enumerate(lines):
        lower = line.lower().strip()
        if any(name in lower for name in section_names):
            capture = True
            continue
        if capture:
            # Stop at next section header (all-caps line or common headers)
            if line.strip() and line.strip().upper() == line.strip() and len(line.strip()) > 3:
                break
            if any(
                kw in lower
                for kw in ["education", "skills", "experience", "project", "certification", "summary", "objective", "reference"]
                if kw not in section_names
            ):
                break
            captured.append(line)
            if len(captured) > 20:
                break
    return "\n".join(captured).strip()


# ---------------------------------------------------------------------------
# Main parse function
# ---------------------------------------------------------------------------

def parse_resume(filepath: str) -> dict:
    """
    Parse a resume file and return a dict with all extracted fields.
    """
    raw_text = extract_text(filepath)
    return {
        "raw_text": raw_text,
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "skills": ", ".join(extract_skills(raw_text)),
        "education": extract_education(raw_text),
        "experience": extract_experience(raw_text),
        "certifications": extract_certifications(raw_text),
        "location": extract_location(raw_text),
    }
